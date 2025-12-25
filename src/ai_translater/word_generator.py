"""Word文档生成模块 - 生成双语对照Word文档"""

from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .pdf_generator import BilingualContent


class WordGenerator:
    """Word文档生成器 - 生成双语对照Word文档"""

    def __init__(
        self,
        font_size: int = 11,
        line_spacing: float = 1.5,
    ):
        """初始化Word生成器
        
        Args:
            font_size: 字体大小
            line_spacing: 行间距
        """
        self.font_size = font_size
        self.line_spacing = line_spacing

    def generate_dual_column_docx(
        self,
        contents: List[BilingualContent],
        output_path: str | Path,
        title: Optional[str] = None,
    ) -> None:
        """生成左右双栏对照Word文档
        
        Args:
            contents: 双语内容列表
            output_path: 输出文件路径
            title: 文档标题
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        
        # 设置默认字体
        self._set_default_font(doc)
        
        # 添加标题
        if title:
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 按页组织内容
        current_page = -1
        for content in contents:
            if content.page_num != current_page:
                current_page = content.page_num
                # 添加页码标题
                page_heading = doc.add_heading(f"第 {current_page + 1} 页", level=2)
                page_heading.runs[0].font.color.rgb = RGBColor(102, 102, 102)

            # 创建双栏表格
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            
            # 设置表格宽度
            for cell in table.rows[0].cells:
                cell.width = Inches(3.0)
            
            # 原文（左列）
            cell_original = table.rows[0].cells[0]
            self._add_text_to_cell(cell_original, content.original, is_original=True)
            
            # 译文（右列）
            cell_translated = table.rows[0].cells[1]
            self._add_text_to_cell(cell_translated, content.translated, is_original=False)
            
            # 添加间距
            doc.add_paragraph()

        # 保存文档
        doc.save(str(output_path))

    def generate_interleaved_docx(
        self,
        contents: List[BilingualContent],
        output_path: str | Path,
        title: Optional[str] = None,
    ) -> None:
        """生成上下交替对照Word文档
        
        Args:
            contents: 双语内容列表
            output_path: 输出文件路径
            title: 文档标题
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._set_default_font(doc)
        
        if title:
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        current_page = -1
        for content in contents:
            if content.page_num != current_page:
                current_page = content.page_num
                page_heading = doc.add_heading(f"第 {current_page + 1} 页", level=2)
                page_heading.runs[0].font.color.rgb = RGBColor(102, 102, 102)

            # 原文（带背景色）
            if content.original:
                para_orig = doc.add_paragraph()
                run = para_orig.add_run(content.original)
                run.font.size = Pt(self.font_size)
                run.font.name = 'Times New Roman'
                # 设置背景色（通过shading）
                self._set_paragraph_shading(para_orig, "F5F5F5")
            
            # 译文
            if content.translated:
                para_trans = doc.add_paragraph()
                run = para_trans.add_run(content.translated)
                run.font.size = Pt(self.font_size)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.color.rgb = RGBColor(51, 51, 51)
            
            doc.add_paragraph()

        doc.save(str(output_path))

    def generate_translation_only_docx(
        self,
        contents: List[BilingualContent],
        output_path: str | Path,
        title: Optional[str] = None,
    ) -> None:
        """生成仅译文Word文档
        
        Args:
            contents: 双语内容列表
            output_path: 输出文件路径
            title: 文档标题
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._set_default_font(doc)
        
        if title:
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        current_page = -1
        for content in contents:
            if content.page_num != current_page:
                current_page = content.page_num
                page_heading = doc.add_heading(f"第 {current_page + 1} 页", level=2)
                page_heading.runs[0].font.color.rgb = RGBColor(102, 102, 102)

            if content.translated:
                para = doc.add_paragraph()
                run = para.add_run(content.translated)
                run.font.size = Pt(self.font_size)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            doc.add_paragraph()

        doc.save(str(output_path))

    def _set_default_font(self, doc: Document) -> None:
        """设置文档默认字体"""
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(self.font_size)
        # 设置中文字体
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def _add_text_to_cell(self, cell, text: str, is_original: bool = True) -> None:
        """向表格单元格添加文本
        
        Args:
            cell: 表格单元格
            text: 文本内容
            is_original: 是否为原文
        """
        # 清除默认段落
        cell.paragraphs[0].clear()
        
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.font.size = Pt(self.font_size)
        
        if is_original:
            run.font.name = 'Times New Roman'
        else:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.color.rgb = RGBColor(51, 51, 51)

    def _set_paragraph_shading(self, para, color: str) -> None:
        """设置段落背景色
        
        Args:
            para: 段落对象
            color: 颜色值（不含#）
        """
        from docx.oxml import OxmlElement
        
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        para._p.get_or_add_pPr().append(shading)

